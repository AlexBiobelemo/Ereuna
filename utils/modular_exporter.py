"""
Modular Exporter Module

This module handles exporting research content into modular documents.
It supports:
1. Individual volume exports (DOCX/PDF)
2. Master document compilation
3. Table of contents generation
4. Cross-references between volumes

Author: Ereuna Development Team
"""

import os
import logging
import zipfile
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from bs4 import BeautifulSoup
import markdown

from utils.docx_generator import DocxGenerator

logger = logging.getLogger(__name__)


@dataclass
class ExportOptions:
    """Options for modular export."""
    output_format: str = "docx"  # 'docx', 'pdf', 'both'
    include_toc: bool = True
    include_cover_page: bool = True
    include_references: bool = True
    cross_reference_volumes: bool = True
    page_size: str = "A4"  # 'A4', 'Letter', 'Legal'
    margins: Tuple[float, float, float, float] = (1.0, 1.0, 1.0, 1.0)  # inches
    font_family: str = "Times New Roman"
    font_size: int = 12
    line_spacing: float = 1.5
    volume_prefix: str = "Volume"


@dataclass
class VolumeExport:
    """Represents an exported volume."""
    volume_number: int
    volume_title: str
    file_path: str
    file_size: int
    pages: int
    sections: List[str]


class ModularExporter:
    """
    Handles modular export of research content.
    
    This class provides functionality to:
    - Export individual volumes as separate documents
    - Generate a master document combining all volumes
    - Create comprehensive table of contents
    - Manage cross-references between volumes
    """
    
    # Constants for styling
    DEFAULT_BATCH_SIZE = 10
    
    def __init__(
        self,
        topic: str,
        output_dir: str = "exports",
        export_options: Optional[ExportOptions] = None
    ):
        """
        Initialize ModularExporter.
        
        Args:
            topic: Research topic
            output_dir: Directory for output files
            export_options: ExportOptions instance (uses defaults if None)
        """
        self.topic = topic
        self.output_dir = output_dir
        self.export_options = export_options or ExportOptions()
        
        # Ensure output directory exists
        self._ensure_output_dir()
        
        # Initialize base DOCX generator
        self.docx_generator = DocxGenerator(topic=topic)
        
        # Track exports
        self.exported_volumes: List[VolumeExport] = []
        self.master_doc_path: Optional[str] = None
    
    def _ensure_output_dir(self):
        """Create output directory if it doesn't exist."""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)
            logger.info(f"Created output directory: {self.output_dir}")
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for filesystem compatibility."""
        # Replace problematic characters
        for char in ['/', '\\', ':', '*', '?', '"', '<', '>', '|']:
            filename = filename.replace(char, '_')
        # Limit length
        if len(filename) > 100:
            filename = filename[:97] + "..."
        return filename
    
    def _create_cover_page(self, document: Document, volume_number: Optional[int] = None):
        """Create a cover page for the document."""
        # Add title
        title = document.add_heading(self.topic, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format title
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add subtitle
        if volume_number:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.add_run(f"{self.export_options.volume_prefix} {volume_number}")
            for run in subtitle.runs:
                run.font.size = Pt(18)
                run.font.italic = True
        else:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.add_run("Comprehensive Research Report")
            for run in subtitle.runs:
                run.font.size = Pt(18)
                run.font.italic = True
        
        # Add date
        date_para = document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        date_para.paragraph_format.space_before = Pt(48)
        
        # Add page break
        document.add_page_break()
    
    def _create_toc_page(self, document: Document, volume_number: Optional[int] = None, toc_data: Optional[List[Dict]] = None):
        """Create a table of contents page."""
        # Add TOC heading
        toc_heading = document.add_heading("Table of Contents", level=1)
        
        if toc_data:
            # Add entries from TOC data
            for entry in toc_data:
                level = entry.get("level", 1)
                title = entry.get("title", "")
                page = entry.get("page", "")
                
                para = document.add_paragraph()
                para.paragraph_format.left_indent = Inches((level - 1) * 0.5)
                para.add_run(title)
                
                # Add dots leader
                if page:
                    para.add_run(f" ....... {page}")
        else:
            # Generic TOC placeholder
            toc_para = document.add_paragraph()
            toc_para.add_run("Volume-specific table of contents will appear here.")
        
        # Add page break
        document.add_page_break()
    
    def _create_references_section(self, document: Document, volume_number: int = 0):
        """Create references section for a volume."""
        # Add References heading
        ref_heading = document.add_heading("References", level=1)
        
        # Add generic reference note
        ref_para = document.add_paragraph()
        ref_para.add_run("References for this volume are included in the master references section at the end of the complete research document.")
        
        if self.export_options.cross_reference_volumes and volume_number > 1:
            # Add cross-reference to Volume 1
            ref_para = document.add_paragraph()
            ref_para.add_run(f"See Volume 1 for the complete bibliography and all cited sources.")
            ref_para.paragraph_format.space_before = Pt(12)
    
    def _convert_markdown_to_docx(self, document: Document, markdown_text: str):
        """Convert markdown content to DOCX, adding to existing document."""
        # Use the existing docx_generator's method
        self.docx_generator._add_markdown_content(document, markdown_text)
    
    def export_single_volume(
        self,
        volume_number: int,
        volume_title: str,
        sections_content: Dict[str, str],
        include_toc: bool = True,
        include_references: bool = True
    ) -> VolumeExport:
        """
        Export a single volume as a DOCX file.
        
        Args:
            volume_number: Volume number
            volume_title: Volume title
            sections_content: Dictionary of section titles to content
            include_toc: Whether to include table of contents
            include_references: Whether to include references section
            
        Returns:
            VolumeExport object with export details
        """
        document = Document()
        
        # Set up document properties
        document.core_properties.title = f"{self.topic} - {volume_title}"
        document.core_properties.author = "Ereuna Research Generator"
        document.core_properties.created = datetime.now()
        
        # Add cover page
        if self.export_options.include_cover_page:
            self._create_cover_page(document, volume_number)
        
        # Generate TOC data
        toc_data = [
            {"title": title, "level": 1, "page": ""}
            for title in sections_content.keys()
        ]
        
        # Add TOC
        if include_toc and self.export_options.include_toc:
            self._create_toc_page(document, volume_number, toc_data)
        
        # Add sections
        for title, content in sections_content.items():
            # Add section heading
            heading = document.add_heading(title, level=1)
            
            # Add content
            self._convert_markdown_to_docx(document, content)
            
            # Add page break between major sections
            if self.export_options.page_size == "A4":
                document.add_page_break()
        
        # Add references section
        if include_references and self.export_options.include_references:
            self._create_references_section(document, volume_number)
        
        # Save document
        safe_topic = self._sanitize_filename(self.topic)
        safe_title = self._sanitize_filename(volume_title)
        filename = f"{safe_topic}_{self.export_options.volume_prefix}_{volume_number}_{safe_title}.docx"
        file_path = os.path.join(self.output_dir, filename)
        
        document.save(file_path)
        
        # Calculate approximate page count
        total_words = sum(len(content.split()) for content in sections_content.values())
        pages = max(1, total_words // 250)  # ~250 words per page
        
        export_info = VolumeExport(
            volume_number=volume_number,
            volume_title=volume_title,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            pages=pages,
            sections=list(sections_content.keys())
        )
        
        self.exported_volumes.append(export_info)
        logger.info(f"Exported Volume {volume_number}: {file_path}")
        
        return export_info
    
    def export_all_volumes(
        self,
        volumes_data: Dict[int, Dict[str, str]],
        volume_titles: Optional[Dict[int, str]] = None
    ) -> List[VolumeExport]:
        """
        Export all volumes as separate documents.
        
        Args:
            volumes_data: Dictionary mapping volume numbers to section content
            volume_titles: Optional dictionary mapping volume numbers to titles
            
        Returns:
            List of VolumeExport objects
        """
        exported = []
        
        for vol_num, sections_content in sorted(volumes_data.items()):
            volume_title = volume_titles.get(vol_num, f"{self.export_options.volume_prefix} {vol_num}")
            export = self.export_single_volume(vol_num, volume_title, sections_content)
            exported.append(export)
        
        logger.info(f"Exported {len(exported)} volumes")
        return exported
    
    def create_master_document(
        self,
        volumes_data: Dict[int, Dict[str, str]],
        volume_titles: Optional[Dict[int, str]] = None,
        master_outline: Optional[List[Dict]] = None,
        generate_toc: bool = True
    ) -> str:
        """
        Create a master document combining all volumes.
        
        Args:
            volumes_data: Dictionary mapping volume numbers to section content
            volume_titles: Optional dictionary mapping volume numbers to titles
            master_outline: Optional master outline structure
            generate_toc: Whether to generate comprehensive TOC
            
        Returns:
            Path to master document
        """
        document = Document()
        
        # Set document properties
        document.core_properties.title = f"{self.topic} - Complete Research Report"
        document.core_properties.author = "Ereuna Research Generator"
        document.core_properties.subject = f"Comprehensive research on {self.topic}"
        
        # Add main cover page
        if self.export_options.include_cover_page:
            self._create_cover_page(document)
        
        # Generate master TOC
        if generate_toc and self.export_options.include_toc:
            toc_data = []
            
            # Add volume entries
            for vol_num in sorted(volumes_data.keys()):
                vol_title = volume_titles.get(vol_num, f"Volume {vol_num}")
                toc_data.append({
                    "title": vol_title,
                    "level": 0,
                    "page": ""
                })
                
                # Add section entries for this volume
                for section_title in volumes_data[vol_num].keys():
                    toc_data.append({
                        "title": section_title,
                        "level": 1,
                        "page": ""
                    })
            
            self._create_toc_page(document, toc_data=toc_data)
        
        # Add each volume
        for vol_num in sorted(volumes_data.keys()):
            vol_title = volume_titles.get(vol_num, f"Volume {vol_num}")
            
            # Add volume header
            vol_heading = document.add_heading(f"{self.export_options.volume_prefix} {vol_num}: {vol_title}", level=0)
            vol_heading.paragraph_format.space_before = Pt(24)
            vol_heading.paragraph_format.space_after = Pt(18)
            
            # Add volume separator
            if self.export_options.cross_reference_volumes:
                vol_note = document.add_paragraph()
                vol_note.add_run(f"This is {vol_title}. Cross-references to other volumes are noted throughout.")
                vol_note.paragraph_format.font.italic = True
                vol_note.paragraph_format.space_after = Pt(12)
            
            # Add sections for this volume
            for section_title, content in volumes_data[vol_num].items():
                # Add section heading
                heading = document.add_heading(section_title, level=1)
                heading.paragraph_format.space_before = Pt(18)
                heading.paragraph_format.space_after = Pt(12)
                
                # Add content
                self._convert_markdown_to_docx(document, content)
                
                # Add page break
                document.add_page_break()
        
        # Save master document
        safe_topic = self._sanitize_filename(self.topic)
        master_filename = f"{safe_topic}_Complete_Research_Report.docx"
        master_path = os.path.join(self.output_dir, master_filename)
        
        document.save(master_path)
        self.master_doc_path = master_path
        
        logger.info(f"Created master document: {master_path}")
        return master_path
    
    def create_volume_manifest(self, output_path: Optional[str] = None) -> str:
        """
        Create a manifest file listing all exported volumes.
        
        Args:
            output_path: Optional custom path for manifest
            
        Returns:
            Path to manifest file
        """
        if not output_path:
            safe_topic = self._sanitize_filename(self.topic)
            output_path = os.path.join(self.output_dir, f"{safe_topic}_manifest.json")
        
        manifest = {
            "topic": self.topic,
            "export_date": datetime.now().isoformat(),
            "total_volumes": len(self.exported_volumes),
            "master_document": self.master_doc_path,
            "volumes": [
                {
                    "volume_number": vol.volume_number,
                    "volume_title": vol.volume_title,
                    "file_path": vol.file_path,
                    "file_size_bytes": vol.file_size,
                    "estimated_pages": vol.pages,
                    "sections": vol.sections
                }
                for vol in self.exported_volumes
            ]
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Created manifest: {output_path}")
        return output_path
    
    def create_zip_archive(self, archive_name: Optional[str] = None) -> str:
        """
        Create a ZIP archive containing all exported volumes.
        
        Args:
            archive_name: Optional custom archive name
            
        Returns:
            Path to ZIP archive
        """
        if not archive_name:
            safe_topic = self._sanitize_filename(self.topic)
            archive_name = f"{safe_topic}_Research_Export.zip"
        
        archive_path = os.path.join(self.output_dir, archive_name)
        
        with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add manifest
            manifest_path = self.create_volume_manifest()
            zipf.write(manifest_path, "manifest.json")
            
            # Add master document
            if self.master_doc_path and os.path.exists(self.master_doc_path):
                zipf.write(
                    self.master_doc_path,
                    os.path.basename(self.master_doc_path)
                )
            
            # Add each volume
            for vol in self.exported_volumes:
                if os.path.exists(vol.file_path):
                    zipf.write(
                        vol.file_path,
                        os.path.basename(vol.file_path)
                    )
        
        logger.info(f"Created ZIP archive: {archive_path}")
        return archive_path
    
    def get_export_summary(self) -> Dict[str, Any]:
        """
        Get summary of all exports.
        
        Returns:
            Dictionary with export summary
        """
        total_size = sum(v.file_size for v in self.exported_volumes)
        total_pages = sum(v.pages for v in self.exported_volumes)
        total_sections = sum(len(v.sections) for v in self.exported_volumes)
        
        return {
            "topic": self.topic,
            "export_date": datetime.now().isoformat(),
            "total_volumes": len(self.exported_volumes),
            "total_estimated_pages": total_pages,
            "total_sections": total_sections,
            "total_file_size_bytes": total_size,
            "total_file_size_mb": round(total_size / (1024 * 1024), 2),
            "master_document": self.master_doc_path,
            "volumes": [
                {
                    "volume_number": v.volume_number,
                    "volume_title": v.volume_title,
                    "file_path": v.file_path,
                    "pages": v.pages,
                    "sections": len(v.sections)
                }
                for v in self.exported_volumes
            ]
        }


# Import json for manifest creation
import json


def create_modular_exporter(
    topic: str,
    output_dir: str = "exports",
    output_format: str = "docx",
    include_toc: bool = True,
    include_cover_page: bool = True,
    cross_reference_volumes: bool = True
) -> ModularExporter:
    """
    Create and return a ModularExporter instance.
    
    Args:
        topic: Research topic
        output_dir: Output directory for exports
        output_format: Export format ('docx', 'pdf', 'both')
        include_toc: Whether to include table of contents
        include_cover_page: Whether to include cover pages
        cross_reference_volumes: Whether to include cross-references
        
    Returns:
        Configured ModularExporter instance
    """
    options = ExportOptions(
        output_format=output_format,
        include_toc=include_toc,
        include_cover_page=include_cover_page,
        cross_reference_volumes=cross_reference_volumes
    )
    
    return ModularExporter(
        topic=topic,
        output_dir=output_dir,
        export_options=options
    )
