from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, Generator, Tuple
import markdown
from bs4 import BeautifulSoup
from bs4.element import Tag, NavigableString
from docx.table import _Cell
from docx.text.paragraph import Paragraph
import logging

logger = logging.getLogger(__name__)

# Constants for styling
HYPERLINK_COLOR = "0000FF"
HYPERLINK_UNDERLINE_STYLE = "single"
CODE_FONT_NAME = "Courier New"
CODE_FONT_SIZE = Pt(10)
TABLE_STYLE = "Table Grid"
DEFAULT_BATCH_SIZE = 10
DEFAULT_HEADING_LEVEL = 3
MIN_HEADING_LEVEL = 1
MAX_HEADING_LEVEL = 9

class DocxGenerator:
    def __init__(self, topic: str, batch_size: int = DEFAULT_BATCH_SIZE):
        self.topic = topic
        self.batch_size = batch_size  # Number of elements to process before yielding

    def _stream_markdown_content(self, markdown_text: str) -> Generator[Tuple[str, Any], None, None]:
        """
        Streams parsed markdown elements as (element_type, element) tuples.
        This allows for memory-efficient processing of large documents.
        """
        html = markdown.markdown(markdown_text, extensions=['fenced_code', 'tables', 'nl2br'])
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.children:
            if element.name:
                yield ('tag', element)
            elif str(element).strip():
                yield ('text', str(element).strip())

    def _add_markdown_content(self, document, markdown_text):
        """
        Parses markdown text and adds it to the document with enhanced styling.
        Uses streaming approach for memory efficiency with large documents.
        """
        table_count = 0
        
        for element_type, element in self._stream_markdown_content(markdown_text):
            if element_type == 'tag' and element.name:
                if element.name.startswith('h') and len(element.name) >= 2:
                    try:
                        level = int(element.name[1])
                        document.add_heading(element.get_text(), level=min(level, MAX_HEADING_LEVEL))
                    except ValueError:
                        document.add_heading(element.get_text(), level=DEFAULT_HEADING_LEVEL)
                elif element.name == 'p':
                    p = document.add_paragraph()
                    self._apply_inline_styles(p, element)
                    # Remove any line numbers that might be in the text
                    if p.text and p.text.startswith('1. '):
                        p.text = p.text[3:]  # Remove "1. " prefix
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        p = document.add_paragraph(style='List Bullet')
                        self._apply_inline_styles(p, li)
                        # Remove any line numbers from list items
                        if p.text and p.text.startswith('1. '):
                            p.text = p.text[3:]
                elif element.name == 'ol':
                    for i, li in enumerate(element.find_all('li')):
                        p = document.add_paragraph(style='List Number')
                        self._apply_inline_styles(p, li)
                        # Remove any line numbers from list items
                        if p.text and p.text.startswith('1. '):
                            p.text = p.text[3:]
                elif element.name == 'blockquote':
                    p = document.add_paragraph(style='Intense Quote')
                    self._apply_inline_styles(p, element)
                elif element.name == 'pre':
                    code_text = element.get_text()
                    p = document.add_paragraph(code_text)
                    for run in p.runs:
                        run.font.name = CODE_FONT_NAME
                        run.font.size = CODE_FONT_SIZE
                elif element.name == 'hr':
                    document.add_paragraph("---", style='Normal')
                elif element.name == 'table':
                    # Process table with batched row handling for memory efficiency
                    table_count += 1
                    self._add_html_table_to_docx(document, element, batch_size=self.batch_size)
                elif element.name == 'img':
                    img_src = element.get('src', 'No source')
                    img_alt = element.get('alt', 'Image')
                    document.add_paragraph(f"[[Image: {img_alt} - {img_src}]]", style='Normal')
                elif element.name == 'a':
                    pass
                else:
                    if element.get_text(strip=True):
                        text = element.get_text()
                        # Remove any line numbers from plain text
                        if text.startswith('1. '):
                            text = text[3:]
                        document.add_paragraph(text)
            elif element_type == 'text':
                text = element
                # Remove any line numbers from plain text
                if text.startswith('1. '):
                    text = text[3:]
                document.add_paragraph(text)
    
    def _add_markdown_content_no_headings(self, document, markdown_text):
        """
        Parses markdown text and adds it to the document with enhanced styling,
        but skips heading processing to avoid duplication.
        Uses streaming approach for memory efficiency with large documents.
        """
        table_count = 0
        
        for element_type, element in self._stream_markdown_content(markdown_text):
            if element_type == 'tag' and element.name:
                # Skip heading elements to avoid duplication
                if element.name.startswith('h') and len(element.name) >= 2:
                    continue
                elif element.name == 'p':
                    p = document.add_paragraph()
                    self._apply_inline_styles(p, element)
                    # Remove any line numbers that might be in the text
                    if p.text and p.text.startswith('1. '):
                        p.text = p.text[3:]  # Remove "1. " prefix
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        p = document.add_paragraph(style='List Bullet')
                        self._apply_inline_styles(p, li)
                        # Remove any line numbers from list items
                        if p.text and p.text.startswith('1. '):
                            p.text = p.text[3:]
                elif element.name == 'ol':
                    for i, li in enumerate(element.find_all('li')):
                        p = document.add_paragraph(style='List Number')
                        self._apply_inline_styles(p, li)
                        # Remove any line numbers from list items
                        if p.text and p.text.startswith('1. '):
                            p.text = p.text[3:]
                elif element.name == 'blockquote':
                    p = document.add_paragraph(style='Intense Quote')
                    self._apply_inline_styles(p, element)
                elif element.name == 'pre':
                    code_text = element.get_text()
                    p = document.add_paragraph(code_text)
                    for run in p.runs:
                        run.font.name = CODE_FONT_NAME
                        run.font.size = CODE_FONT_SIZE
                elif element.name == 'hr':
                    document.add_paragraph("---", style='Normal')
                elif element.name == 'table':
                    # Process table with batched row handling for memory efficiency
                    table_count += 1
                    self._add_html_table_to_docx(document, element, batch_size=self.batch_size)
                elif element.name == 'img':
                    img_src = element.get('src', 'No source')
                    img_alt = element.get('alt', 'Image')
                    document.add_paragraph(f"[[Image: {img_alt} - {img_src}]]", style='Normal')
                elif element.name == 'a':
                    pass
                else:
                    if element.get_text(strip=True):
                        text = element.get_text()
                        # Remove any line numbers from plain text
                        if text.startswith('1. '):
                            text = text[3:]
                        document.add_paragraph(text)
            elif element_type == 'text':
                text = element
                # Remove any line numbers from plain text
                if text.startswith('1. '):
                    text = text[3:]
                document.add_paragraph(text)

    def _apply_inline_styles(self, paragraph: Paragraph, soup_element: Any):
        """
        Applies inline styles (bold, italic, inline code, hyperlinks) to a paragraph
        based on the parsed BeautifulSoup element's contents.
        """
        # Set default paragraph formatting for better appearance
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        
        for content in soup_element.contents:
            if content.name == 'strong':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'code': # Inline code
                run = paragraph.add_run(content.get_text())
                run.font.name = CODE_FONT_NAME # Apply monospace font for inline code
                run.font.size = CODE_FONT_SIZE
            elif content.name == 'a': # Hyperlink
                # Add a hyperlink to the paragraph
                self._add_hyperlink(paragraph, content.get_text(), content.get('href', '#'))
            elif content.string:
                # Add plain text content
                paragraph.add_run(content.string)
            elif content.name: # Handle nested tags like <strong><em>text</em></strong>
                # Recursively apply styles for nested elements
                self._apply_inline_styles(paragraph, content)

    def _add_hyperlink(self, paragraph, text, url):
        """
        Adds a hyperlink to a paragraph.
        This involves creating a relationship to the external URL and
        embedding the hyperlink XML element into the paragraph.
        """
        part = paragraph.part
        # Create a relationship to the external URL
        rId = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # Create the w:hyperlink XML element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), rId,) # Set the relationship ID
        
        # Create a new run for the hyperlink text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr') # Run properties
        new_run.append(rPr)
        
        # Add color and underline to hyperlink text for visual indication
        c = OxmlElement('w:color')
        c.set(qn('w:val'), HYPERLINK_COLOR) # Blue color
        rPr.append(c)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), HYPERLINK_UNDERLINE_STYLE) # Single underline
        rPr.append(u)
        
        # Add the text content to the run
        new_run.add_child_element(OxmlElement('w:t', text=text))
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink) # Append the hyperlink to the paragraph's XML
        return hyperlink

    def _add_html_table_to_docx(self, document, html_table, batch_size: int = 10):
        """
        Converts an HTML table (BeautifulSoup object) into a DOCX table.
        Uses batched row processing for memory efficiency with large tables.
        
        Args:
            document: The DOCX document to add the table to.
            html_table: The BeautifulSoup element containing the HTML table.
            batch_size: Number of rows to process before yielding (for streaming).
        """
        # Extract table headers
        headers = [th.get_text(strip=True) for th in html_table.find_all('th')]
        
        # Collect all rows first (for determining column count)
        all_rows = []
        for tr in html_table.find_all('tr'):
            cols = [td.get_text(strip=True) for td in tr.find_all('td')]
            if cols:
                all_rows.append(cols)

        if not headers and not all_rows:
            return

        # Determine the number of columns based on headers or the first row
        num_cols = len(headers) if headers else (len(all_rows[0]) if all_rows else 0)
        if num_cols == 0:
            return

        # Add a new table to the document with an initial row
        table = document.add_table(rows=1, cols=num_cols)
        table.style = TABLE_STYLE

        # Add headers to the first row of the DOCX table
        if headers:
            hdr_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                if i < num_cols:
                    hdr_cells[i].text = header_text
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Add data rows in batches for memory efficiency
        start_row_idx = 1 if headers else 0
        data_rows = all_rows[start_row_idx:]
        
        for i in range(0, len(data_rows), batch_size):
            batch = data_rows[i:i + batch_size]
            for row_data in batch:
                # Ensure row_data has enough elements for the columns
                row_data_padded = row_data + [''] * (num_cols - len(row_data))
                cells = table.add_row().cells
                for j, cell_text in enumerate(row_data_padded):
                    if j < num_cols:
                        cells[j].text = cell_text

    def generate_docx_report(self, sections_content: Dict[str, str], output_path: str):
        """
        Generates a DOCX report from a dictionary of section titles and their markdown content.
        """
        document = Document()

        # Add the main topic as a level 1 heading
        document.add_heading(self.topic, level=1)

        # Iterate through each section and add its title and content
        for title, content in sections_content.items():
            # Format the title with the main topic if needed
            formatted_title = title.replace("{topic}", self.topic)
            document.add_heading(formatted_title, level=2) # Add section title as a level 2 heading
            self._add_markdown_content(document, content) # Add markdown content for the section

        document.save(output_path)
